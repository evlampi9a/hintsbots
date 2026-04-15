"""
Hints.ai Telegram Bot
- Answers questions about past KP using RAG from Supabase
- Transcribes voice messages to text (Whisper)
- Generates cost estimates using the Hints pricing prompt (now with Claude Sonnet)
- Reads attached files (PDF, DOCX, TXT) to enrich context
- Groups multiple messages/files sent in quick succession (debounce)
- Formats responses for Telegram readability (no tables, max 2 nesting levels)
- Query expansion: LLM generates synonyms/related terms before vector search
- Keyword fallback: if vector search finds too few results, adds keyword search
"""
import os, requests, tempfile, logging, asyncio, re
from telegram import Update
from telegram.ext import ApplicationBuilder, CommandHandler, MessageHandler, filters, ContextTypes
from analytical_handler import is_analytical_query, handle_analytical_query
from similar_projects_handler import is_similar_query, handle_similar_query
from openai import OpenAI
import anthropic

logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(message)s")
logger = logging.getLogger(__name__)

# ── Config ───────────────────────────────────────────────────────────────────
TELEGRAM_TOKEN = os.environ["TELEGRAM_TOKEN"]
OPENAI_KEY = os.environ["OPENAI_KEY"]
ANTHROPIC_KEY = os.environ["ANTHROPIC_KEY"]
SB_URL = os.environ["SB_URL"]
SB_KEY = os.environ["SB_KEY"]
SB_H = {"apikey": SB_KEY, "Authorization": f"Bearer {SB_KEY}", "Content-Type": "application/json"}

DEBOUNCE_SECONDS = 8  # Увеличен чтобы успеть поймать текст после файла
SEARCH_COUNT = 30

with open(os.path.join(os.path.dirname(__file__), "pricing_prompt.txt")) as f:
    PRICING_PROMPT = f.read()

oai = OpenAI(api_key=OPENAI_KEY, base_url="https://api.openai.com/v1")
ant = anthropic.Anthropic(api_key=ANTHROPIC_KEY)

# ── Conversation history helpers ─────────────────────────────────────────────
HISTORY_MAX = 20  # keep last 20 turns (user+assistant pairs)

def is_new_topic(new_text: str, history: list[dict]) -> bool:
    """
    Ask a fast LLM whether the new message starts a NEW project/topic
    compared to the existing conversation history.
    Returns True if topic has changed (history should be reset).
    """
    if not history:
        return False
    # Build a compact summary of the last few turns for the LLM
    recent = history[-6:]  # last 3 pairs max
    history_snippet = "\n".join(h["role"] + ": " + h["content"][:300] for h in recent)
    try:
        resp = oai.chat.completions.create(
            model="gpt-4.1-nano",
            messages=[
                {
                    "role": "system",
                    "content": (
                        "Ты определяешь, является ли новое сообщение пользователя продолжением "
                        "текущего диалога или началом нового, несвязанного проекта/темы.\n"
                        "Отвечай ТОЛЬКО одним словом: CONTINUE или NEW.\n\n"
                        "CONTINUE — если новое сообщение:\n"
                        "- уточняет, дополняет или развивает текущую тему\n"
                        "- просит пересчитать/переделать то же самое с другим параметром\n"
                        "- задаёт вопрос по тому же проекту/компании/отрасли\n"
                        "- говорит 'сделай расчёт', 'найди похожие', 'а если онлайн?' и т.п.\n\n"
                        "NEW — если новое сообщение:\n"
                        "- описывает совершенно другой проект, другую компанию, другую отрасль\n"
                        "- явно меняет тему (например, с отелей на банки, с B2B на B2C)\n"
                        "- начинает с нуля без связи с предыдущим\n\n"
                        "При сомнении выбирай CONTINUE."
                    ),
                },
                {
                    "role": "user",
                    "content": f"ИСТОРИЯ ДИАЛОГА:\n{history_snippet}\n\nНОВОЕ СООБЩЕНИЕ: {new_text}",
                },
            ],
            temperature=0.0,
            max_tokens=5,
        )
        answer = resp.choices[0].message.content.strip().upper()
        logger.info(f"Topic detection: '{answer}' for: {new_text[:60]}")
        return answer == "NEW"
    except Exception as e:
        logger.error(f"Topic detection error: {e}")
        return False  # on error — keep context, safer


def update_history(context: ContextTypes.DEFAULT_TYPE, role: str, content: str):
    """Append a message to history, keeping last HISTORY_MAX entries."""
    history = context.user_data.get("history", [])
    history.append({"role": role, "content": content})
    context.user_data["history"] = history[-HISTORY_MAX:]


def get_history_context(context: ContextTypes.DEFAULT_TYPE) -> str:
    """Return formatted history string for injection into system prompt (excluding last user msg)."""
    history = context.user_data.get("history", [])
    if len(history) <= 1:
        return ""
    return "\n".join(h["role"] + ": " + h["content"][:500] for h in history[:-1])


# ── Query expansion ───────────────────────────────────────────────────────────
def expand_query(query: str) -> tuple[list[str], list[str]]:
    """
    Use a fast LLM to generate:
    - 4 expanded query variants for vector search
    - a list of specific keywords/terms for keyword search (countries, cities, brands, etc.)
    Returns (queries_list, keywords_list).
    """
    try:
        resp = oai.chat.completions.create(
            model="gpt-4.1-nano",
            messages=[
                {
                    "role": "system",
                    "content": (
                        "Ты помогаешь расширить поисковый запрос для базы коммерческих предложений "
                        "исследовательского агентства. Верни JSON объект с двумя полями:\n"
                        "- \"queries\": список из 4 вариантов запроса (синонимы, смежные термины, английские эквиваленты)\n"
                        "- \"keywords\": список конкретных слов для поиска по тексту: "
                        "названия стран, городов, компаний, брендов, аббревиатуры, специфические термины. "
                        "Например для 'арабские страны': [\"ОАЭ\", \"UAE\", \"Dubai\", \"Дубай\", \"Саудовская\", \"KSA\", \"Saudi\", \"Катар\", \"Qatar\", \"Бахрейн\", \"MENA\", \"GCC\", \"арабск\"]\n"
                        "Верни ТОЛЬКО валидный JSON, без пояснений."
                    ),
                },
                {"role": "user", "content": f"Запрос: {query}"},
            ],
            temperature=0.2,
            max_tokens=400,
            response_format={"type": "json_object"},
        )
        import json
        data = json.loads(resp.choices[0].message.content)
        queries = data.get("queries", [])
        keywords = data.get("keywords", [])
        # Combine original + expanded queries, deduplicate
        all_queries = [query] + [q.strip() for q in queries if q.strip()]
        seen, result = set(), []
        for q in all_queries:
            if q and q not in seen:
                seen.add(q)
                result.append(q)
        logger.info(f"Query expansion queries: {result}")
        logger.info(f"Query expansion keywords: {keywords}")
        return result[:5], keywords[:20]
    except Exception as e:
        logger.error(f"Query expansion error: {e}")
        # Fallback: extract words from query as keywords
        words = re.findall(r'\b\w{3,}\b', query, re.UNICODE)
        return [query], words[:10]

# ── RAG helpers ───────────────────────────────────────────────────────────────
def embed(text: str):
    return oai.embeddings.create(
        model="text-embedding-3-small", input=text[:8000]
    ).data[0].embedding

def search_kp_single(query: str, threshold: float = 0.28, count: int = SEARCH_COUNT):
    """Single vector search for one query string."""
    emb = embed(query)
    r = requests.post(
        f"{SB_URL}/rest/v1/rpc/match_kp_chunks",
        headers=SB_H,
        json={"query_embedding": emb, "match_threshold": threshold, "match_count": count},
        timeout=20,
    )
    if r.status_code == 200:
        return r.json()
    logger.error(f"Supabase error {r.status_code}: {r.text[:200]}")
    return []

def keyword_search_kp(keywords: list[str], limit: int = 200) -> list[dict]:
    """
    Search kp_chunks by keyword ILIKE match using a single OR query per field.
    Returns rows with project_name, company, chunk_text (no similarity score).
    """
    kws = [kw for kw in keywords if len(kw) >= 3]
    if not kws:
        return []
    results = {}
    # Build a single OR filter: chunk_text=ilike.*kw1*,ilike.*kw2*,...
    for field in ["chunk_text", "project_name"]:
        or_filter = ",".join(f"ilike.*{kw}*" for kw in kws)
        r = requests.get(
            f"{SB_URL}/rest/v1/kp_chunks",
            headers={k: v for k, v in SB_H.items() if k != "Content-Type"},
            params={
                "select": "project_name,company,chunk_text,sheet_name",
                field: or_filter,
                "limit": limit,
            },
            timeout=15,
        )
        if r.status_code == 200:
            for row in r.json():
                pname = row.get("project_name", "")
                if pname and pname not in results:
                    results[pname] = {
                        "project_name": pname,
                        "company": row.get("company", ""),
                        "chunk_text": row.get("chunk_text", ""),
                        "sheet_name": row.get("sheet_name", ""),
                        "similarity": 0.0,
                    }
    return list(results.values())

def get_company_for_project(project_name: str) -> str:
    """Look up company from КП or metadata sheet for a given project name."""
    r = requests.get(
        f"{SB_URL}/rest/v1/kp_chunks",
        headers={k: v for k, v in SB_H.items() if k != "Content-Type"},
        params={
            "select": "company",
            "project_name": f"eq.{project_name}",
            "company": "neq.",
            "sheet_name": "in.(КП,metadata)",
            "limit": 1,
        },
        timeout=10,
    )
    if r.status_code == 200:
        rows = r.json()
        if rows and rows[0].get("company"):
            return rows[0]["company"]
    return ""

# company cache to avoid repeated lookups
_company_cache: dict[str, str] = {}

def resolve_company(project_name: str, company: str) -> str:
    """Return company if already set, otherwise look it up from КП/metadata."""
    if company:
        return company
    if project_name in _company_cache:
        return _company_cache[project_name]
    resolved = get_company_for_project(project_name)
    _company_cache[project_name] = resolved
    return resolved

def search_kp(query: str, threshold: float = 0.35, count: int = SEARCH_COUNT):
    """
    Full search pipeline:
    1. Expand query into variants + keywords via LLM
    2. ALWAYS run keyword search first (exact/partial matches)
    3. Run vector search for each variant with higher threshold (0.35)
    4. Merge: keyword results first (most precise), then vector-only results
    5. Resolve company for all results
    """
    queries, kw_list = expand_query(query)

    # Step 1: ALWAYS run keyword search (finds exact name matches like "Айбим ОАЭ")
    keyword_projects: dict[str, dict] = {}
    if kw_list:
        kw_results = keyword_search_kp(kw_list, limit=100)
        for res in kw_results:
            pname = res.get("project_name", "")
            if pname:
                keyword_projects[pname] = res
        logger.info(f"Keyword search found: {len(keyword_projects)} projects")

    # Step 2: vector search across all expanded queries (higher threshold = less noise)
    vector_projects: dict[str, dict] = {}
    for q in queries:
        results = search_kp_single(q, threshold=threshold, count=count)
        for res in results:
            pname = res.get("project_name", "")
            if not pname:
                continue
            # Keep the result with highest similarity for each project
            if pname not in vector_projects or res.get("similarity", 0) > vector_projects[pname].get("similarity", 0):
                vector_projects[pname] = res
    logger.info(f"Vector search found: {len(vector_projects)} projects")

    # Step 3: merge — keyword results first (marked as keyword), then vector-only
    seen_projects: dict[str, dict] = {}
    # Add keyword results first
    for pname, res in keyword_projects.items():
        seen_projects[pname] = res
    # Add vector results that weren't found by keyword
    for pname, res in vector_projects.items():
        if pname not in seen_projects:
            seen_projects[pname] = res
        else:
            # If already in keyword results, update similarity score
            seen_projects[pname]["similarity"] = res.get("similarity", 0)

    # Step 4: batch-resolve company for all projects missing it
    missing_company = [pname for pname, res in seen_projects.items() if not res.get("company")]
    if missing_company:
        # Single request to get company for all missing projects
        or_filter = ",".join(f"eq.{pname}" for pname in missing_company)
        r = requests.get(
            f"{SB_URL}/rest/v1/kp_chunks",
            headers={k: v for k, v in SB_H.items() if k != "Content-Type"},
            params={
                "select": "project_name,company",
                "project_name": or_filter,
                "company": "neq.",
                "sheet_name": "in.(КП,metadata)",
                "limit": 500,
            },
            timeout=15,
        )
        if r.status_code == 200:
            for row in r.json():
                pname = row.get("project_name", "")
                company = row.get("company", "")
                if pname and company and pname in seen_projects:
                    seen_projects[pname]["company"] = company
                    _company_cache[pname] = company
    final = list(seen_projects.values())

    # Sort: by similarity desc (keyword-only results have 0.0, appear last)
    final.sort(key=lambda x: x.get("similarity", 0), reverse=True)
    return final

def build_kp_context(results: list) -> str:
    parts = []
    for r in results:
        name = r['project_name']
        company = r['company'] or "—"
        sim = r.get('similarity', 0)
        chunk = r['chunk_text']
        sim_str = f"{sim:.0%}" if sim > 0 else "keyword match"
        parts.append(
            f"Проект: {name} | Компания: {company} "
            f"(релевантность: {sim_str})\n{chunk}"
        )
    return "\n\n---\n\n".join(parts)

def build_sources_text(results: list) -> str:
    seen, sources = set(), []
    for r in results:
        k = r['project_name']
        if k not in seen:
            seen.add(k)
            name = r['project_name'].replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            company = (r['company'] or "—").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            sim = r.get('similarity', 0)
            sim_str = f"{sim:.0%}" if sim > 0 else "по ключевым словам"
            sources.append(f"• <b>{name}</b> ({company}) — {sim_str}")
    return "\n".join(sources)

# ── Telegram formatting ───────────────────────────────────────────────────────
ALLOWED_HTML_TAGS = re.compile(
    r"<(/?)(?:b|i|u|s|code|pre|a)(\s[^>]*)?>|<br\s*/?>|&(?:amp|lt|gt|quot|apos);",
    re.IGNORECASE
)

def table_to_text(match_obj) -> str:
    """Конвертирует Markdown-таблицу в читаемый текст."""
    lines = [l for l in match_obj.group(0).strip().split("\n") if l.strip()]
    if len(lines) < 3:
        return match_obj.group(0)
    headers = [h.strip() for h in lines[0].strip("|").split("|")]
    result_rows = []
    for line in lines[2:]:
        if "|" not in line:
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        row_parts = []
        for i, cell in enumerate(cells):
            header = headers[i] if i < len(headers) else ""
            if header and cell:
                row_parts.append(f"<b>{header}:</b> {cell}")
            elif cell:
                row_parts.append(cell)
        if row_parts:
            result_rows.append("  ".join(row_parts))
    return "\n".join(result_rows)

def post_process_for_telegram(text: str) -> str:
    """Конвертирует Markdown/смешанный ответ LLM в чистый HTML для Telegram."""
    # 1. Конвертируем Markdown-таблицы
    text = re.sub(r"(\|[^\n]+\|\n)([\|\-: ]+\|\n)(\|[^\n]+\|\n)+", table_to_text, text)
    # 2. Конвертируем Markdown → HTML (до экранирования, чтобы не сломать уже вставленные теги)
    # ## Заголовок → <b>Заголовок</b>
    text = re.sub(r"^#{1,6}\s+(.+)$", lambda m: f"<b>{m.group(1).strip()}</b>", text, flags=re.MULTILINE)
    # **жирный** → <b>жирный</b>
    text = re.sub(r"\*\*(.+?)\*\*", lambda m: f"<b>{m.group(1)}</b>", text)
    # *жирный* → <b>жирный</b> (одиночные звёздочки, не часть **)
    text = re.sub(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)", lambda m: f"<b>{m.group(1)}</b>", text)
    # _курсив_ → <i>курсив</i>
    text = re.sub(r"(?<!_)_(?!_)(.+?)(?<!_)_(?!_)", lambda m: f"<i>{m.group(1)}</i>", text)
    # `код` → <code>код</code>
    text = re.sub(r"`([^`]+)`", lambda m: f"<code>{m.group(1)}</code>", text)
    # 3. Убираем разделители ---
    text = re.sub(r"^---+$", "", text, flags=re.MULTILINE)
    # 4. Вложенные списки (4+ пробела) → подпункт
    text = re.sub(r"^(\s{4,})[\-•]\s+", r"    ▸ ", text, flags=re.MULTILINE)
    # 5. Маркеры списков → • (звёздочка, дефис, точка в начале строки)
    text = re.sub(r"^\s*[\*\-•]\ +", "• ", text, flags=re.MULTILINE)
    # 6. Экранируем & которые не являются HTML-сущностями
    text = re.sub(r"&(?!amp;|lt;|gt;|quot;|apos;)", "&amp;", text)
    # 7. Убираем лишние пустые строки
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()

# ── File reading ──────────────────────────────────────────────────────────────
def extract_text_from_file(path: str, ext: str) -> str:
    ext = ext.lower().lstrip(".")
    if ext == "pdf":
        import pypdf
        reader = pypdf.PdfReader(path)
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    if ext == "docx":
        import docx as _docx
        doc = _docx.Document(path)
        return "\n".join(p.text for p in doc.paragraphs)
    if ext in ("txt", "md", "csv"):
        with open(path, "r", errors="replace") as f:
            return f.read()
    return ""

# ── Intent detection ──────────────────────────────────────────────────────────
PRICING_KEYWORDS = [
    "посчитай", "расчет", "рассчитай", "сколько стоит", "смета", "бюджет",
    "стоимость", "цена", "просчёт", "просчет", "прайс", "оцени проект",
    "сделай расчет", "сделай просчет", "коммерческое предложение", "кп на",
    "калькуляция", "сколько будет стоить",
]
CREATIVE_KEYWORDS = [
    # Текстовые задачи
    "напиши", "сформулируй", "составь", "помоги написать", "напишем",
    "ответ клиенту", "письмо", "сообщение клиенту", "питч", "текст",
    "перефразируй", "переформулируй", "как написать", "как ответить",
    "как сказать", "как объяснить", "помоги с текстом", "отправить клиенту",
    "подготовь ответ", "сделай ответ", "одним сообщением", "одним письмом",
    # Документы и файлы
    "презентация", "презентацию", "слайды", "декс", "документ", "док",
    "таблицу", "таблица", "ексель", "табличку", "спредшит",
    "бриф", "брифинг", "резюме", "отчёт", "отчет", "сводку", "свод", "докс",
    "ворд", "пдф", "файл", "скачать", "скачай",
]

def is_pricing_request(text: str) -> bool:
    t = text.lower()
    return any(kw in t for kw in PRICING_KEYWORDS)

def is_creative_request(text: str) -> bool:
    t = text.lower()
    return any(kw in t for kw in CREATIVE_KEYWORDS)

def has_pricing_in_history(history: list[dict]) -> bool:
    """Проверяет, есть ли уже расчёт в текущем диалоге."""
    pricing_markers = [
        "Трудозатраты", "Цена для клиента", "Себестоимость", "Прямые затраты",
        "Тимлид:", "Ресёрчер:", "Стандартный просчёт", "Итого цена", "Итого себестоимость",
    ]
    for h in history:
        if h["role"] == "assistant":
            if any(marker in h["content"] for marker in pricing_markers):
                return True
    return False

def classify_intent(text: str, history: list[dict]) -> str:
    """
    LLM-классификация намерения пользователя с учётом контекста диалога.
    Возвращает одно из: PRICING | CREATIVE | ANALYTICAL | SIMILAR | RAG
    """
    recent = history[-6:] if history else []
    history_snippet = "\n".join(h["role"] + ": " + h["content"][:400] for h in recent) if recent else "(нет истории)"
    try:
        resp = oai.chat.completions.create(
            model="gpt-4.1-nano",
            messages=[
                {
                    "role": "system",
                    "content": (
                        "Ты классифицируешь намерение пользователя в боте агентства исследований Hints.\n"
                        "Учитывай историю диалога — пользователь может продолжать предыдущую тему.\n\n"
                        "Категории:\n\n"
                        "PRICING — пользователь хочет получить расчёт стоимости исследования, смету, бюджет, цену.\n"
                        "  Признаки: любое упоминание цены/стоимости/бюджета, просьба посчитать/оценить/просчитать проект,\n"
                        "  запрос на КП, уточнение цифр из уже сделанного расчёта, пересчёт с другими параметрами.\n"
                        "  Примеры: 'посчитай 20 глубинок', 'сколько будет стоить', 'а если сделать онлайн?', 'дай цену без рекрутинга'.\n\n"
                        "CREATIVE — пользователь хочет что-то создать или написать: любой текст, файл, документ.\n"
                        "  Признаки: просьба написать/составить/сформулировать/перефразировать, упоминание письма/питча/\n"
                        "  презентации/документа/таблицы/брифа/отчёта, запрос ответить клиенту одним сообщением.\n"
                        "  Примеры: 'напиши ответ клиенту', 'сделай презентацию', 'составь таблицу', 'помоги ответить', 'подготовь бриф'.\n\n"
                        "ANALYTICAL — пользователь хочет цифры, статистику или агрегированные данные по базе проектов.\n"
                        "  Признаки: вопрос о средней/медианной цене, диапазоне цен, частоте типов проектов,\n"
                        "  топе сфер, популярных работах, распределении чего-либо по базе.\n"
                        "  Примеры: 'медианная цена FinTech', 'стоимость B2B проектов в среднем', 'как часто делаем глубинки', 'диапазон цен'.\n\n"
                        "SIMILAR — пользователь хочет найти/посмотреть проекты с фильтрацией по параметрам.\n"
                        "  Признаки: запрос найти/показать/отфильтровать проекты, упоминание бюджета/сферы/типа работ\n"
                        "  в контексте поиска проектов из базы.\n"
                        "  Примеры: 'найди B2B до 600к', 'покажи проекты в FinTech', 'есть ли глубинки до 300к', 'покажи все проекты Сбера'.\n\n"
                        "RAG — общий вопрос по опыту/компетенциям, вопрос о конкретных проектах, что делали для компании.\n"
                        "  Примеры: 'есть ли опыт в MENA', 'что делали для Яндекса', 'как рекрутируем топ-менеджеров'.\n\n"
                        "Отвечай ТОЛЬКО одним словом: PRICING, CREATIVE, ANALYTICAL, SIMILAR, RAG"
                    ),
                },
                {
                    "role": "user",
                    "content": f"ИСТОРИЯ ДИАЛОГА:\n{history_snippet}\n\nНОВЫЙ ЗАПРОС: {text}",
                },
            ],
            temperature=0.0,
            max_tokens=10,
        )
        intent = resp.choices[0].message.content.strip().upper()
        if intent not in ("PRICING", "CREATIVE", "ANALYTICAL", "SIMILAR", "RAG"):
            intent = "RAG"
        logger.info(f"Intent classified: {intent} for: {text[:60]}")
        return intent
    except Exception as e:
        logger.error(f"Intent classification error: {e} — fallback to RAG")
        return "RAG"  # Безопасный фаллбэк: RAG универсален, не угадываем по словам

# ── Telegram send ─────────────────────────────────────────────────────────────
async def send_long(update: Update, text: str):
    MAX = 4000
    chunks = [text[i : i + MAX] for i in range(0, len(text), MAX)]
    for chunk in chunks:
        try:
            await update.message.reply_text(chunk, parse_mode="HTML")
        except Exception as e:
            logger.warning(f"HTML parse failed: {e}, sending plain")
            # Убираем HTML-теги и отправляем как plain text
            plain = re.sub(r"<[^>]+>", "", chunk)
            plain = plain.replace("&amp;", "&").replace("&lt;", "<").replace("&gt;", ">")
            await update.message.reply_text(plain)

# ── Debounce ──────────────────────────────────────────────────────────────────
async def _flush_and_process(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await asyncio.sleep(DEBOUNCE_SECONDS)
    # Если есть файлы но нет текста — ждём ещё немного (пользователь мог отправить файл раньше текста)
    if context.user_data.get("pending_files") and not context.user_data.get("pending_texts"):
        logger.info("Files without text — waiting extra 6s for user text...")
        await asyncio.sleep(6)
    texts = context.user_data.pop("pending_texts", [])
    files = context.user_data.pop("pending_files", [])
    context.user_data["debounce_task"] = None

    combined_text = "\n".join(t for t in texts if t)
    file_sections = []
    file_names = []
    for (fname, ext, path) in files:
        try:
            content = extract_text_from_file(path, ext)
            if content.strip():
                # Truncate file content for history (keep first 1500 chars)
                file_sections.append(f"=== Файл: {fname} ===\n{content.strip()}")
                file_names.append(fname)
        except Exception as e:
            logger.error(f"File read error ({fname}): {e}")
            file_sections.append(f"=== Файл: {fname} === [не удалось прочитать]")
    if file_sections:
        combined_text += "\n\n" + "\n\n".join(file_sections)
    combined_text = combined_text.strip()
    if not combined_text:
        return

    # Save combined message (text + files) to history for context
    # Build a compact history entry: user text + file names (not full content — too long)
    history_entry_parts = []
    if texts:
        history_entry_parts.append("\n".join(t for t in texts if t))
    if file_names:
        history_entry_parts.append("Приложены файлы: " + ", ".join(file_names))
    history_entry = "\n".join(history_entry_parts)

    # Detect topic change before updating history
    history = context.user_data.get("history", [])
    if history and is_new_topic(history_entry, history):
        logger.info(f"New topic detected in flush — resetting history")
        context.user_data["history"] = []
    if history_entry:
        update_history(context, "user", history_entry)

    user_text = "\n".join(t for t in texts if t).strip()
    await process_query(update, context, combined_text, user_text=user_text)

def _schedule_flush(update: Update, context: ContextTypes.DEFAULT_TYPE):
    old = context.user_data.get("debounce_task")
    if old and not old.done():
        old.cancel()
    task = asyncio.create_task(_flush_and_process(update, context))
    context.user_data["debounce_task"] = task

# ── Core processing ───────────────────────────────────────────────────────────
async def process_query(update: Update, context: ContextTypes.DEFAULT_TYPE, text: str, user_text: str = ""):
    history = context.user_data.get("history", [])
    # Для классификации используем оригинальный запрос пользователя + краткий контекст файла
    # чтобы текст файла не «перебивал» намерение
    classify_text = text
    has_file = len(text) > len(user_text) + 50 if user_text else "=== Файл:" in text
    if user_text and has_file:
        # Есть файл + текст пользователя — передаём запрос + первые 200 символов файла как подсказку
        file_preview = text[len(user_text):].strip()[:200]
        classify_text = f"{user_text}\n[Приложен файл, начало: {file_preview}]"
    elif has_file and not user_text:
        # Файл без текста — это почти всегда бриф для расчёта
        logger.info("File without user text — forcing PRICING intent")
        intent = "PRICING"
        logger.info(f"Routing to handler: PRICING (forced for file)")
        await handle_pricing(update, context, text)
        return
    intent = classify_intent(classify_text, history)
    logger.info(f"Routing to handler: {intent}")
    if intent == "ANALYTICAL":
        await handle_analytical_query(update, context, text)
    elif intent == "SIMILAR":
        await handle_similar_query(update, context, text)
    elif intent == "PRICING":
        await handle_pricing(update, context, text)
    elif intent == "CREATIVE":
        await handle_creative(update, context, text)
    else:
        await handle_rag(update, context, text)

async def handle_pricing(update: Update, context: ContextTypes.DEFAULT_TYPE, text: str):
    history = context.user_data.get("history", [])
    is_followup = has_pricing_in_history(history)

    if is_followup:
        await update.message.reply_text("⏳ Пересчитываю...")
    else:
        await update.message.reply_text("⏳ Делаю расчёт (Claude Sonnet 4.5)...")

    try:
        kp_results = search_kp(text, threshold=0.25, count=6)
        kp_context = build_kp_context(kp_results)
        history_context = get_history_context(context)

        if is_followup:
            # Гибкий режим: пользователь уточняет или пересчитывает уже сделанный расчёт
            system_content = (
                "Ты — коммерческий ассистент агентства Hints. Уже есть сделанный расчёт в истории диалога.\n"
                "Пользователь уточняет, пересчитывает или задаёт вопрос по уже сделанному расчёту.\n\n"
                "ПРАВИЛА:\n"
                "- НЕ повторяй весь расчёт целиком. Отвечай только на то, что спросили.\n"
                "- Если просят пересчитать с другими параметрами — дай только изменившиеся цифры (трудозатраты, цена, сроки).\n"
                "- Если задают вопрос по методологии — ответь конкретно и без шаблона.\n"
                "- ФОРМАТИРОВАНИЕ: только HTML-теги Telegram (<b>, <i>). Без звёздочек, решёток, таблиц.\n"
                "- Слово 'маржа' не использовать.\n"
            )
            if history_context:
                system_content += f"\nИСТОРИЯ ДИАЛОГА:\n{history_context}"
            if kp_context:
                system_content += "\n\nПохожие проекты для калибровки:\n" + kp_context
            user_msg = text
        else:
            # Первый расчёт: полный шаблон из 12 блоков
            system_content = PRICING_PROMPT
            if history_context:
                system_content += f"\n\nИСТОРИЯ ДИАЛОГА (используй для понимания контекста — не повторяй уже сказанное):\n{history_context}"
            if kp_context:
                system_content += (
                    "\n\nДополнительный контекст — похожие проекты из базы КП Hints "
                    "(используй для калибровки, не копируй структуру):\n\n" + kp_context
                )
            user_msg = f"Сделай стандартный просчёт для следующего запроса:\n\n{text}"

        completion = ant.messages.create(
            model="claude-sonnet-4-5",
            system=system_content,
            messages=[
                {"role": "user", "content": user_msg},
            ],
            temperature=0.2,
            max_tokens=4000,
        )
        raw_response = completion.content[0].text
        update_history(context, "assistant", raw_response)
        response = post_process_for_telegram(raw_response)
        if kp_results and not is_followup:
            response += "\n\n📎 <b>Похожие проекты в базе:</b>\n" + build_sources_text(kp_results[:5])
        await send_long(update, response)
    except Exception as e:
        logger.error(f"Pricing error: {e}")
        await update.message.reply_text("Ошибка при расчёте. Попробуй ещё раз.")

async def handle_creative(update: Update, context: ContextTypes.DEFAULT_TYPE, text: str):
    """Обрабатывает задачи на создание контента: письма, презентации, документы, таблицы, брифы."""
    import json as _json
    import tempfile as _tempfile
    import os as _os
    try:
        history_context = get_history_context(context)
        kp_results = search_kp(text, threshold=0.3, count=5)
        kp_context = build_kp_context(kp_results) if kp_results else ""

        # Шаг 1: определяем формат вывода через лёгкий LLM
        format_resp = oai.chat.completions.create(
            model="gpt-4.1-nano",
            messages=[
                {
                    "role": "system",
                    "content": (
                        "Определи формат вывода для задачи. Отвечай ТОЛЬКО одним словом:\n"
                        "DOCX — если просят документ, ворд, док, бриф, отчёт, сводку\n"
                        "XLSX — если просят таблицу, ексель, спредшит, сводную таблицу\n"
                        "TEXT — если просят письмо, сообщение, питч, ответ, текст для отправки"
                    ),
                },
                {"role": "user", "content": text},
            ],
            temperature=0.0,
            max_tokens=5,
        )
        output_format = format_resp.choices[0].message.content.strip().upper()
        if output_format not in ("DOCX", "XLSX", "TEXT"):
            output_format = "TEXT"
        logger.info(f"Creative output format: {output_format}")

        # Шаг 2: строим системный промпт
        hints_context = (
            "Ты — коммерческий ассистент агентства исследований Hints.\n\n"
            "О Hints:\n"
            "- Клиенты: крупный бизнес (Яндекс, Сбер, Авито, Пятёрочка, Самолёт, Т-Банк и др.)\n"
            "- Специализация: сложные B2B-исследования, труднодоступная ЦА, зарубежные рынки (MENA, Азия, СНГ, Европа, США)\n"
            "- Услуги: качественные (глубинные интервью, фокус-группы), количественные (опросы), "
            "кабинетные исследования, рекрутинг, экспертные интервью, UX-исследования\n"
            "- Сильные стороны: экспертиза в труднодоступной аудитории, международный опыт, "
            "прозрачная методология, работа с топ-клиентами\n"
        )

        if output_format == "TEXT":
            await update.message.reply_text("✍️ Пишу...")
            system_msg = (
                hints_context +
                "ПРАВИЛА:\n"
                "1. Пиши живо и по-человечески — без канцелярита и шаблонных фраз\n"
                "2. Используй контекст из истории диалога — цифры, детали проекта, расчёты\n"
                "3. Если просят письмо/сообщение — выдай готовый текст, который можно скопировать и отправить\n"
                "4. Если просят питч — структурируй кратко: проблема → решение → почему мы → следующий шаг\n"
                "5. Тон: профессиональный, но живой. Не официозный\n"
                "6. ФОРМАТИРОВАНИЕ ДЛЯ TELEGRAM: <b>жирный</b> для ключевого, списки через дефис, без таблиц\n"
            )
            if history_context:
                system_msg += f"\nКОНТЕКСТ ДИАЛОГА (цифры и детали):\n{history_context}"
            if kp_context:
                system_msg += f"\nПОХОЖИЕ ПРОЕКТЫ (примеры для сильных сторон):\n{kp_context}"
            completion = ant.messages.create(
                model="claude-sonnet-4-5",
                system=system_msg,
                messages=[{"role": "user", "content": text}],
                temperature=0.6,
                max_tokens=3000,
            )
            raw_response = completion.content[0].text
            update_history(context, "assistant", raw_response)
            response = post_process_for_telegram(raw_response)
            await send_long(update, response)

        elif output_format == "DOCX":
            await update.message.reply_text("✍️ Готовлю документ...")
            system_msg = (
                hints_context +
                "Создай структурированный документ в формате Markdown.\n"
                "Используй # для заголовков, ## для подзаголовков, - для списков, **жирный** для выделения.\n"
                "Используй контекст из истории диалога — цифры, детали, расчёты.\n"
            )
            if history_context:
                system_msg += f"КОНТЕКСТ ДИАЛОГА:\n{history_context}\n"
            if kp_context:
                system_msg += f"ПОХОЖИЕ ПРОЕКТЫ:\n{kp_context}\n"
            completion = ant.messages.create(
                model="claude-sonnet-4-5",
                system=system_msg,
                messages=[{"role": "user", "content": text}],
                temperature=0.4,
                max_tokens=4000,
            )
            md_content = completion.content[0].text
            update_history(context, "assistant", md_content[:500])
            # Создаём DOCX
            from docx import Document as _Document
            from docx.shared import Pt as _Pt
            doc = _Document()
            for line in md_content.split("\n"):
                if line.startswith("# "):
                    p = doc.add_heading(line[2:], level=1)
                elif line.startswith("## "):
                    p = doc.add_heading(line[3:], level=2)
                elif line.startswith("### "):
                    p = doc.add_heading(line[4:], level=3)
                elif line.startswith("- ") or line.startswith("* "):
                    doc.add_paragraph(line[2:], style="List Bullet")
                elif line.strip():
                    p = doc.add_paragraph()
                    # Обработка **жирного**
                    parts = re.split(r'\*\*(.+?)\*\*', line)
                    for i, part in enumerate(parts):
                        run = p.add_run(part)
                        if i % 2 == 1:
                            run.bold = True
            tmp = _tempfile.NamedTemporaryFile(suffix=".docx", delete=False, dir="/tmp")
            doc.save(tmp.name)
            with open(tmp.name, "rb") as f:
                await update.message.reply_document(document=f, filename="hints_document.docx")
            _os.unlink(tmp.name)

        elif output_format == "XLSX":
            await update.message.reply_text("📊 Готовлю таблицу...")
            system_msg = (
                hints_context +
                "Создай таблицу в формате JSON. Верни ТОЛЬКО JSON без пояснений.\n"
                "Формат: {\"title\": \"Название таблицы\", \"headers\": [\"Колонка 1\", ...], \"rows\": [[\"значение\", ...], ...]}\n"
                "Используй контекст из истории диалога для заполнения данными.\n"
            )
            if history_context:
                system_msg += f"КОНТЕКСТ ДИАЛОГА:\n{history_context}\n"
            completion = ant.messages.create(
                model="claude-sonnet-4-5",
                system=system_msg,
                messages=[{"role": "user", "content": text}],
                temperature=0.3,
                max_tokens=3000,
            )
            raw = completion.content[0].text.strip()
            # Извлекаем JSON если есть лишний текст
            json_match = re.search(r'\{.*\}', raw, re.DOTALL)
            table_data = _json.loads(json_match.group() if json_match else raw)
            update_history(context, "assistant", f"Таблица: {table_data.get('title', '')}, {len(table_data.get('rows', []))} строк")
            import openpyxl as _openpyxl
            from openpyxl.styles import Font as _Font, PatternFill as _PatternFill
            wb = _openpyxl.Workbook()
            ws = wb.active
            ws.title = table_data.get("title", "Hints")[:31]
            headers = table_data.get("headers", [])
            for col, h in enumerate(headers, 1):
                cell = ws.cell(row=1, column=col, value=h)
                cell.font = _Font(bold=True)
                cell.fill = _PatternFill("solid", fgColor="E8F0FE")
            for row_i, row in enumerate(table_data.get("rows", []), 2):
                for col_i, val in enumerate(row, 1):
                    ws.cell(row=row_i, column=col_i, value=val)
            # Автоширина колонок
            for col in ws.columns:
                max_len = max((len(str(c.value or "")) for c in col), default=10)
                ws.column_dimensions[col[0].column_letter].width = min(max_len + 4, 50)
            tmp = _tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False, dir="/tmp")
            wb.save(tmp.name)
            with open(tmp.name, "rb") as f:
                await update.message.reply_document(document=f, filename="hints_table.xlsx")
            _os.unlink(tmp.name)

    except Exception as e:
        logger.error(f"Creative error: {e}")
        await update.message.reply_text("Ошибка при генерации. Попробуй ещё раз.")

async def handle_rag(update: Update, context: ContextTypes.DEFAULT_TYPE, text: str):
    await update.message.reply_text("🔍 Ищу по всей базе КП (Claude Sonnet 4.5)...")
    try:
        results = search_kp(text, threshold=0.25, count=SEARCH_COUNT)
        if not results:
            await update.message.reply_text(
                "По вашему запросу ничего не найдено в базе КП. "
                "Попробуйте переформулировать или задать другой вопрос."
            )
            return
        context_text = build_kp_context(results)
        system_msg = (
            "Ты — ассистент агентства исследований Hints. "
            "Отвечай на вопросы по базе коммерческих предложений агентства.\n\n"
            "ПРАВИЛА ФОРМАТИРОВАНИЯ ДЛЯ TELEGRAM (СТРОГО ОБЯЗАТЕЛЬНО):\n"
            "1. ЗАПРЕЩЕНО: таблицы, разделители ---, Markdown-символы (* _ # ` []).\n"
            "2. ЖИРНЫЙ: пиши <b>текст</b> — только для заголовков секций и ключевых данных.\n"
            "3. КУРСИВ: пиши <i>текст</i> — для пояснений.\n"
            "4. СПИСКИ: только дефис + пробел: - пункт. Вложенный: 4 пробела + - подпункт.\n"
            "5. ЭМОДЗИ: используй перед заголовками секций (не в отдельных строках-заголовках).\n"
            "   Пример: 💡 <b>Что мы можем сделать:</b>\n"
            "6. Пустые строки между блоками обязательны.\n\n"
            "ПРИМЕР ПРАВИЛЬНОГО ОТВЕТА:\n"
            "Да, у нас есть опыт рекрутинга такой аудитории.\n\n"
            "💡 <b>Что мы можем сделать:</b>\n"
            "- Найти 8-9 сотрудников международных отелей\n"
            "- Которые работают с экстранетами Booking, Expedia\n\n"
            "📋 <b>Похожие проекты:</b>\n"
            "- Дроздов отели — работа с гостиничным бизнесом\n"
            "- Сбер TaxFree — рекрут сотрудников отелей\n\n"
            "❓ <b>Что нужно уточнить:</b>\n"
            "- География отелей (Россия, СНГ, другие страны?)\n"
            "- Формат встречи (онлайн/офлайн)\n\n"
            "ТЕБЕ ПЕРЕДАЮТСЯ ФРАГМЕНТЫ ПРОЕКТОВ ИЗ БАЗЫ КП. "
            "Твоя задача: прочитай все фрагменты и ответь на вопрос пользователя. "
            "ФИЛЬТРАЦИЯ: если проект упоминает страну/регион лишь как бенчмарк — не включай его. "
            "Если проект реально проводился в этой стране (аудитория там, респонденты там) — включай. "
            "ВСЕ релевантные проекты перечисляй — не ограничивайся 3-5 примерами. "
            "Для каждого релевантного проекта укажи: название, компанию, страну/регион, краткую суть. "
            "Отвечай по-русски, структурированно. Без вступлений и комплиментов.\n\n"
            "УТОЧНЯЮЩИЕ ВОПРОСЫ: если запрос неясный, неполный или не хватает данных для ответа — ЗАДАЙ УТОЧНЯЮЩИЕ ВОПРОСЫ. "
            "Не пытайся угадать. Лучше спроси, чем дай неточный ответ. "
            "Пример: если просят \"проекты в Азии\", спроси: \"Какие страны Азии вас интересуют? Китай, Индия, Япония, ЮВА?\""
            "\n\nЧТО ДАЛЬШЕ: в конце каждого ответа предлагай 2-3 следующих шага. "
            "Примеры:\n"
            "- «Что ещё можно сделать:»\n"
            "- «Дальше можно:»\n"
            "Предлагай релевантные действия: сделать расчёт, найти похожие проекты, посмотреть аналитику.\n"
            "Пример: «Дальше можно: <b>сделать расчёт</b> для этого проекта или <b>найти похожие</b> B2B проекты в России.»"
        )
        history_context = get_history_context(context)
        if history_context:
            system_msg += f"\n\nИСТОРИЯ ДИАЛОГА (используй для понимания контекста — не повторяй уже сказанное):\n{history_context}"

        user_msg = f"Вопрос: {text}\n\nФРАГМЕНТЫ ПРОЕКТОВ ИЗ БАЗЫ КП:\n\n{context_text}"
        completion = ant.messages.create(
            model="claude-sonnet-4-5",
            system=system_msg,
            messages=[
                {"role": "user", "content": user_msg},
            ],
            temperature=0.1,
            max_tokens=3000,
        )
        raw_response = completion.content[0].text
        update_history(context, "assistant", raw_response)
        response = post_process_for_telegram(raw_response)
        await send_long(update, response)
    except Exception as e:
        logger.error(f"RAG error: {e}")
        await update.message.reply_text("Произошла ошибка. Попробуй ещё раз.")

# ── App setup ──────────────────────────────────────────────────────────────────
async def cmd_start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text(
        "Привет! Я — AI ассистент агентства Hints 👋\n\n"
        "<b>Я умею:</b>\n"
        "• Отвечать на вопросы по базе КП (220+ проектов)\n"
        "• Делать первичный расчёт стоимости по ТЗ или брифу\n"
        "• Читать прикреплённые файлы (PDF, DOCX, TXT)\n"
        "• Принимать голосовые сообщения\n\n"
        "Можешь отправить сразу несколько сообщений и файлов — я подожду и обработаю всё вместе.\n\n"
        "<b>Примеры запросов:</b>\n"
        "— Какие проекты делали для Яндекса?\n"
        "— Дай все КП по ОАЭ / арабским странам\n"
        "— Посчитай: 20 глубинных интервью с топ-менеджерами в MENA\n"
        "— Аналитика по FinTech проектам\n"
        "— Найди B2B проекты в России до 600к",
        parse_mode="HTML",
    )

async def cmd_help(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text(
        "<b>🔍 Вопросы по базе КП:</b>\n"
        "Задай вопрос — бот найдёт все релевантные проекты.\n"
        "<i>«Какие проекты делали в FinTech?»</i>, <i>«Все КП по ОАЭ»</i>\n\n"
        "<b>💰 Расчёт стоимости:</b>\n"
        "Напиши ТЗ или прикрепи бриф (PDF/DOCX/TXT) — бот сделает просчёт.\n"
        "<i>«Посчитай 15 глубинок с HR-директорами»</i>\n\n"
        "<b>📊 Аналитика:</b>\n"
        "<i>«Медианная цена FinTech проектов»</i>, <i>«Статистика по B2B»</i>\n\n"
        "<b>🔎 Поиск похожих:</b>\n"
        "<i>«Найди B2B проекты в России до 600к»</i>, <i>«Проекты Яндекса»</i>\n\n"
        "<b>📎 Несколько сообщений сразу:</b>\n"
        "Отправляй текст и файлы подряд — бот подождёт 3 сек и обработает всё вместе.\n\n"
        "<b>🎙 Голос:</b> Можно отправить голосовое — бот транскрибирует и ответит.",
        parse_mode="HTML",
    )

async def handle_text(update: Update, context: ContextTypes.DEFAULT_TYPE):
    text = update.message.text.strip()
    logger.info(f"Text [{update.effective_user.id}]: {text[:80]}")
    # History is updated in _flush_and_process after debounce (so files are included too)
    context.user_data.setdefault("pending_texts", []).append(text)
    context.user_data.setdefault("pending_files", [])
    _schedule_flush(update, context)

async def handle_voice(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text("🎤 Транскрибирую голосовое...")
    try:
        file = await update.message.voice.get_file()
        with tempfile.NamedTemporaryFile(suffix=".ogg", delete=False) as tmp:
            await file.download_to_drive(tmp.name)
            tmp_path = tmp.name
        with open(tmp_path, "rb") as af:
            transcript = oai.audio.transcriptions.create(model="whisper-1", file=af)
        text = transcript.text
        logger.info(f"Voice transcribed [{update.effective_user.id}]: {text[:80]}")
        await update.message.reply_text(f"📝 Распознал: _{text}_", parse_mode="Markdown")
        # Voice transcript goes into pending_texts — history updated in _flush_and_process
        context.user_data.setdefault("pending_texts", []).append(f"🎤 Голосовое: {text}")
        context.user_data.setdefault("pending_files", [])
        _schedule_flush(update, context)
    except Exception as e:
        logger.error(f"Voice error: {e}")
        await update.message.reply_text("Не удалось обработать голосовое. Попробуй ещё раз.")

async def handle_document(update: Update, context: ContextTypes.DEFAULT_TYPE):
    doc = update.message.document
    fname = doc.file_name or "file"
    ext = fname.rsplit(".", 1)[-1] if "." in fname else ""
    supported = {"pdf", "docx", "txt", "md", "csv"}
    if ext.lower() not in supported:
        await update.message.reply_text(
            f"⚠️ Формат .{ext} не поддерживается. Поддерживаются: PDF, DOCX, TXT."
        )
        return
    await update.message.reply_text(f"📎 Получил файл *{fname}*, добавлю в запрос...", parse_mode="Markdown")
    try:
        tg_file = await doc.get_file()
        with tempfile.NamedTemporaryFile(suffix=f".{ext}", delete=False) as tmp:
            await tg_file.download_to_drive(tmp.name)
            local_path = tmp.name
        # Если к файлу есть подпись (caption) — добавляем её как текст запроса
        caption = update.message.caption
        if caption and caption.strip():
            logger.info(f"File caption detected: {caption[:80]}")
            context.user_data.setdefault("pending_texts", []).append(caption.strip())
        else:
            context.user_data.setdefault("pending_texts", [])
        context.user_data.setdefault("pending_files", []).append((fname, ext, local_path))
        _schedule_flush(update, context)
    except Exception as e:
        logger.error(f"Document error: {e}")
        await update.message.reply_text(f"Не удалось скачать файл {fname}.")

def main():
    app = (
        ApplicationBuilder()
        .token(TELEGRAM_TOKEN)
        .connect_timeout(30)
        .read_timeout(30)
        .write_timeout(60)
        .pool_timeout(30)
        .build()
    )
    app.add_handler(CommandHandler("start", cmd_start))
    app.add_handler(CommandHandler("help", cmd_help))
    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_text))
    app.add_handler(MessageHandler(filters.VOICE, handle_voice))
    app.add_handler(MessageHandler(filters.Document.ALL, handle_document))
    logger.info("Hints KP bot starting...")
    app.run_polling(drop_pending_updates=True)

if __name__ == "__main__":
    main()
